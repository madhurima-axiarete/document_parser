"""
pdfminer_extractor.py

Extracts text and structure from documents using pdfminer.six.
Handles PDFs well; produces empty text for image files (jpg/png).

Parsing logic ported from:
  /Users/madhurimachakraborty/dr/app/services/extraction/parsers.py
"""

from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Optional

METHOD = "pdfminer"

# ── KV patterns ────────────────────────────────────────────────────────────────

_KV_COLON_RE = re.compile(
    r"^(?P<key>[A-Za-z][A-Za-z0-9 /\-\(\)]{1,60}?)\s*[:\-=]\s*(?P<value>.{1,200})\s*$",
    re.MULTILINE,
)
_KV_TABULAR_RE = re.compile(
    r"^(?P<key>[A-Za-z][A-Za-z0-9 /\-\(\)]{2,40}?)\s{2,}(?P<value>\S.{0,200})\s*$",
    re.MULTILINE,
)


# ── Section detection ──────────────────────────────────────────────────────────


def _extract_sections(text: str) -> list[dict]:
    candidates: list[tuple[int, str, int]] = []

    for m in re.finditer(r"^#+\s+(.+)$", text, re.MULTILINE):
        level = len(m.group(0)) - len(m.group(0).lstrip("#"))
        candidates.append((m.start(), m.group(1).strip(), level))

    for m in re.finditer(r"^(\d+(?:\.\d+)*)[\.)]?\s+([A-Z][^\n]{2,79})\s*$", text, re.MULTILINE):
        depth = m.group(1).count(".") + 1
        candidates.append((m.start(), m.group(2).strip(), depth))

    for m in re.finditer(r"^([A-Z][A-Z0-9\s\-\:]{3,79})\s*$", text, re.MULTILINE):
        title = m.group(1).strip()
        if len(title.split()) <= 10 and title == title.upper():
            candidates.append((m.start(), title.title(), 2))

    for m in re.finditer(r"^([A-Z][^\n]{3,79})\n[=\-]{3,}", text, re.MULTILINE):
        level = 1 if "=" in m.group(0).splitlines()[-1] else 2
        candidates.append((m.start(), m.group(1).strip(), level))

    if not candidates:
        return []

    candidates.sort(key=lambda x: x[0])
    seen_offsets: set[int] = set()
    unique: list[tuple[int, str, int]] = []
    for offset, title, level in candidates:
        if not any(abs(offset - s) < 5 for s in seen_offsets):
            unique.append((offset, title, level))
            seen_offsets.add(offset)

    sections = []
    for i, (start, title, level) in enumerate(unique):
        end = unique[i + 1][0] if i + 1 < len(unique) else len(text)
        content = text[start:end].strip()
        lines = content.splitlines()
        body = "\n".join(lines[1:]).strip() if lines else ""
        sections.append({"title": title, "level": level, "content": body[:500]})

    return sections


# ── KVP extraction ─────────────────────────────────────────────────────────────


def _extract_kvps(text: str, sections: list[dict]) -> list[dict]:
    kvps: list[dict] = []
    seen_keys: set[str] = set()

    def _section_at(offset: int) -> Optional[str]:
        sec = None
        char_pos = 0
        for s in sections:
            if char_pos <= offset:
                sec = s["title"]
            char_pos += len(s.get("content", "")) + len(s.get("title", "")) + 10
        return sec

    def _add(key: str, value: str, offset: int, source: str) -> None:
        k, v = key.strip(), value.strip()
        if not k or not v or len(v) > 300:
            return
        if k.lower() in seen_keys:
            return
        seen_keys.add(k.lower())
        kvps.append({"key": k, "value": v, "section": _section_at(offset), "source": source})

    for m in _KV_COLON_RE.finditer(text):
        _add(m.group("key"), m.group("value"), m.start(), "label_colon")
    for m in _KV_TABULAR_RE.finditer(text):
        _add(m.group("key"), m.group("value"), m.start(), "label_tabular")

    return kvps


# ── PDF parsing ────────────────────────────────────────────────────────────────


def _parse_pdf(content: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        from pdfminer.high_level import extract_text
        raw = extract_text(io.BytesIO(content)) or ""
    except Exception as exc:
        warnings.append(f"pdfminer error: {exc}")
        raw = ""
    if not raw.strip():
        warnings.append("PDF produced no text — may be a scanned image")
    return raw, warnings


# ── Table extraction from DOCX ─────────────────────────────────────────────────


def _parse_docx(content: bytes) -> tuple[str, list[dict], list[dict], list[str]]:
    warnings: list[str] = []
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        return "", [], [], [f"python-docx error: {exc}"]

    raw_lines: list[str] = []
    tables: list[dict] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            raw_lines.append(text)

    for idx, tbl in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        rows = [r for r in rows if any(c for c in r)]
        if rows:
            tables.append({
                "headers": rows[0],
                "rows": rows[1:],
                "source": f"table {idx + 1}",
            })

    raw_text = "\n".join(raw_lines)
    sections = _extract_sections(raw_text)
    kvps = _extract_kvps(raw_text, sections)
    return raw_text, sections, tables, kvps, warnings


# ── Public interface ───────────────────────────────────────────────────────────


def extract(file_path: str) -> dict:
    """Extract text and structure from a document using pdfminer/python-docx."""
    path = Path(file_path)
    content = path.read_bytes()
    ext = path.suffix.lower().lstrip(".")
    warnings: list[str] = []

    if ext in ("jpg", "jpeg", "png", "gif", "bmp", "tiff", "webp"):
        warnings.append(f"pdfminer cannot process image files (.{ext})")
        return {
            "file": path.name,
            "method": METHOD,
            "raw_text_chars": 0,
            "sections": [],
            "tables": [],
            "key_value_pairs": [],
            "warnings": warnings,
        }

    if ext == "pdf":
        raw_text, w = _parse_pdf(content)
        warnings.extend(w)
        sections = _extract_sections(raw_text)
        kvps = _extract_kvps(raw_text, sections)
        tables: list[dict] = []

    elif ext == "docx":
        result = _parse_docx(content)
        raw_text, sections, tables, kvps, w = result
        warnings.extend(w)

    elif ext == "csv":
        try:
            text = content.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = [r for r in reader if any(c.strip() for c in r)]
            headers = rows[0] if rows else []
            data_rows = rows[1:]
            raw_text = "\n".join(",".join(r) for r in rows)
            tables = [{"headers": headers, "rows": data_rows, "source": "csv"}]
            sections = []
            kvps = [{"key": r[0].strip(), "value": r[1].strip(), "section": None, "source": "csv"}
                    for r in data_rows if len(r) >= 2 and r[0].strip() and r[1].strip()] \
                    if len(headers) == 2 else []
        except Exception as exc:
            warnings.append(f"CSV parse error: {exc}")
            raw_text, sections, tables, kvps = "", [], [], []

    else:
        try:
            raw_text = content.decode("utf-8", errors="replace")
        except Exception as exc:
            warnings.append(f"Text decode error: {exc}")
            raw_text = ""
        sections = _extract_sections(raw_text)
        kvps = _extract_kvps(raw_text, sections)
        tables = []

    return {
        "file": path.name,
        "method": METHOD,
        "raw_text_chars": len(raw_text),
        "sections": sections,
        "tables": tables,
        "key_value_pairs": kvps,
        "warnings": warnings,
    }
