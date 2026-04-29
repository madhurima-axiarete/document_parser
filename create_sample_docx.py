"""
create_sample_docx.py

Generates a sample DOCX with complex features:
- Headers/footers, complex tables, nested lists, mixed formatting,
  embedded images, footnotes, and headings.

Run once: python create_sample_docx.py
Output: SampleDocument.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io


def _make_png(rgb: tuple, w=200, h=150) -> bytes:
    """Generate a PNG image in memory."""
    img = Image.new("RGB", (w, h), color=rgb)
    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


def _set_cell_background(cell, fill):
    """Set cell background color."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


doc = Document()

# ── Header ─────────────────────────────────────────────────────────────────

section = doc.sections[0]
header = section.header
h_para = header.paragraphs[0]
h_para.text = "CONFIDENTIAL — Sample Document"
h_para.runs[0].font.size = Pt(10)
h_para.runs[0].font.italic = True

# ── Footer ─────────────────────────────────────────────────────────────────

footer = section.footer
f_para = footer.paragraphs[0]
f_para.text = "Page "
run = f_para.add_run()
run.add_break()  # Page number field will be handled differently
f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Cover page ─────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.text = "Sample Document"
title_run = title.runs[0]
title_run.font.size = Pt(36)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.text = "Complex Formatting & Layout Testing"
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(16)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# ── Abstract ───────────────────────────────────────────────────────────────

abstract = doc.add_paragraph()
abstract.text = "Abstract"
abstract.style = "Heading 1"

abstract_body = doc.add_paragraph(
    "This document tests extraction of complex DOCX features including "
    "tables with merged cells, nested tables, embedded images, multiple font styles, "
    "nested lists, headers/footers, and mixed formatting."
)

# ── Introduction ───────────────────────────────────────────────────────────

intro = doc.add_paragraph("Introduction", style="Heading 1")

intro_para = doc.add_paragraph(
    "Document extraction is a challenging task. Complex layouts, "
    "varied formatting, and embedded elements require careful handling."
)

# ── Section 1: Tables ──────────────────────────────────────────────────────

section1 = doc.add_paragraph("Tables and Structure", style="Heading 1")

doc.add_paragraph(
    "Tables can have merged cells, alternating row colors, and nested structure:"
)

# Main table: 6 rows x 4 cols
table = doc.add_table(rows=6, cols=4)
table.style = "Light Grid Accent 1"

# Header row
header_cells = table.rows[0].cells
headers = ["Column 1", "Column 2", "Column 3", "Column 4"]
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    cell.text = header_text
    _set_cell_background(cell, "4A90E2")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows with alternating colors
for row_idx in range(1, 6):
    row = table.rows[row_idx]
    for col_idx, cell in enumerate(row.cells):
        cell.text = f"R{row_idx}C{col_idx + 1}"
        if row_idx % 2 == 0:
            _set_cell_background(cell, "F0F0F0")

# ── Nested Table ───────────────────────────────────────────────────────────

doc.add_paragraph("Nested tables demonstrate hierarchical structure:", style="Heading 2")

outer_table = doc.add_table(rows=2, cols=2)
outer_table.style = "Light Grid Accent 1"

# Top-left cell with nested content
cell = outer_table.rows[0].cells[0]
cell.text = "Cell 1: Parent"

# Top-right cell with nested table
cell = outer_table.rows[0].cells[1]
nested_table = cell.add_table(rows=2, cols=2)
nested_table.style = "Light Grid"
for i, row in enumerate(nested_table.rows):
    for j, ncell in enumerate(row.cells):
        ncell.text = f"N{i + 1}C{j + 1}"

# Bottom cells
outer_table.rows[1].cells[0].text = "Cell 3"
outer_table.rows[1].cells[1].text = "Cell 4"

# ── Section 2: Images ──────────────────────────────────────────────────────

section2 = doc.add_paragraph("Embedded Images", style="Heading 1")

doc.add_paragraph("Three inline images with captions:")

for color_rgb, label in [((220, 100, 100), "Red"), ((100, 220, 100), "Green"), ((100, 100, 220), "Blue")]:
    para = doc.add_paragraph()
    img_bytes = _make_png(color_rgb)
    img_stream = io.BytesIO(img_bytes)
    para.add_run().add_picture(img_stream, width=Inches(1.5))
    para.add_run(f"  {label} image").font.size = Pt(10)

# ── Section 3: Formatting ──────────────────────────────────────────────────

section3 = doc.add_paragraph("Mixed Formatting", style="Heading 1")

# Various text styles
format_para = doc.add_paragraph()
format_para.add_run("Bold text").bold = True
format_para.add_run(" | ")
format_para.add_run("Italic text").italic = True
format_para.add_run(" | ")
colored_run = format_para.add_run("Colored text (red)")
colored_run.font.color.rgb = RGBColor(255, 0, 0)
format_para.add_run(" | ")
size_run = format_para.add_run("Different size (12pt)")
size_run.font.size = Pt(12)

# Strikethrough (via OpenXML)
strikethrough_para = doc.add_paragraph()
strikethrough_run = strikethrough_para.add_run("Strikethrough text")
strikethrough_run.font.strikethrough = True

# ── Section 4: Lists ───────────────────────────────────────────────────────

section4 = doc.add_paragraph("Nested Lists", style="Heading 1")

list_items = [
    ("Level 0 Item 1", 0),
    ("Level 1 Item 1.1", 1),
    ("Level 2 Item 1.1.1", 2),
    ("Level 2 Item 1.1.2", 2),
    ("Level 1 Item 1.2", 1),
    ("Level 0 Item 2", 0),
    ("Level 1 Item 2.1", 1),
    ("Level 2 Item 2.1.1", 2),
]

for text, level in list_items:
    para = doc.add_paragraph(text, style="List Bullet")
    # Adjust paragraph level for indentation
    para.paragraph_format.left_indent = Inches(level * 0.25)

# Numbered list mixed in
doc.add_paragraph("Numbered list:", style="Heading 2")
for i, text in enumerate(["First step", "Second step", "Third step"], 1):
    para = doc.add_paragraph(text, style="List Number")

# ── Section 5: Footnotes ───────────────────────────────────────────────────

section5 = doc.add_paragraph("Footnotes and References", style="Heading 1")

para = doc.add_paragraph("This document contains important information")
footnote_run = para.add_run("[1]")
footnote_run.font.superscript = True
para.add_run(" that requires careful extraction.")

para2 = doc.add_paragraph("Additional details are provided")
footnote_run2 = para2.add_run("[2]")
footnote_run2.font.superscript = True
para2.add_run(" in the appendix.")

doc.add_paragraph()
doc.add_paragraph("[1] This is the first footnote content.", style="Normal")
doc.add_paragraph("[2] This is the second footnote content.", style="Normal")

# ── Section 6: References ──────────────────────────────────────────────────

references = doc.add_paragraph("References", style="Heading 1")
doc.add_paragraph("Smith, J. (2025). Document Parsing. Journal of AI.", style="Normal")
doc.add_paragraph("Johnson, A. (2024). Complex Formats. Tech Press.", style="Normal")

# ── Save ───────────────────────────────────────────────────────────────────

doc.save("SampleDocument.docx")
print("Created SampleDocument.docx")
